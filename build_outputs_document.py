from __future__ import annotations

import argparse
import importlib.util
import json
import math
import sys
import textwrap
from datetime import date
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent
DATA_DIR = PROJECT_ROOT / "data"
DELIVERABLE_DIR = PROJECT_ROOT / "deliverables"
DEFAULT_ASSET_DIR = DELIVERABLE_DIR / "outputs_report_assets"
DEFAULT_OUTPUT_DOCX = DELIVERABLE_DIR / "prediction_market_volatility_outputs_report.docx"
SKILL_ROOT = (
    Path.home()
    / ".codex"
    / "plugins"
    / "cache"
    / "openai-primary-runtime"
    / "documents"
    / "26.521.10419"
    / "skills"
    / "documents"
)


def load_table_geometry_helper():
    helper_path = SKILL_ROOT / "scripts" / "table_geometry.py"
    spec = importlib.util.spec_from_file_location("table_geometry", helper_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load table geometry helper from {helper_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


TABLE_GEOMETRY = load_table_geometry_helper()

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BODY_COLOR = RGBColor(31, 37, 45)
MUTED_COLOR = RGBColor(96, 110, 128)
ACCENT_BLUE = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
GRID_GRAY = "D8DEE9"
HEADER_FILL = "F2F4F7"
CHART_COLORS = [
    "#2E74B5",
    "#7A5A00",
    "#1F4D78",
    "#9B1C1C",
    "#6A3D9A",
    "#00897B",
    "#455A64",
]


def ensure_dirs() -> None:
    DELIVERABLE_DIR.mkdir(parents=True, exist_ok=True)
    DEFAULT_ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_saved_data(data_dir: Path = DATA_DIR) -> dict[str, object]:
    analytical = pd.read_csv(data_dir / "analytical_dataset.csv", parse_dates=["date"])
    scoring = pd.read_csv(data_dir / "scoring_dataset.csv", parse_dates=["date"])
    rankings = pd.read_csv(data_dir / "category_rankings.csv")
    latest_scores = pd.read_csv(data_dir / "latest_market_scores.csv", parse_dates=["date"])
    coefficients = pd.read_csv(data_dir / "ols_coefficients.csv")
    importances = pd.read_csv(data_dir / "rf_feature_importances.csv")
    volatility_over_time = pd.read_csv(data_dir / "volatility_over_time.csv", parse_dates=["date"])
    correlation_matrix = pd.read_csv(data_dir / "correlation_matrix.csv", index_col=0)
    with open(data_dir / "model_metrics.json", "r", encoding="utf-8") as handle:
        metrics = json.load(handle)
    metadata = {
        "excluded_categories": [],
        "excluded_market_count": 0,
        "excluded_observation_count": 0,
        "filter_label": "All categories",
        "uses_saved_summary_artifacts": True,
    }
    metadata_path = data_dir / "filter_metadata.json"
    if metadata_path.exists():
        with open(metadata_path, "r", encoding="utf-8") as handle:
            metadata.update(json.load(handle))
    return {
        "analytical": analytical,
        "scoring": scoring,
        "rankings": rankings,
        "latest_scores": latest_scores,
        "coefficients": coefficients,
        "importances": importances,
        "volatility_over_time": volatility_over_time,
        "correlation_matrix": correlation_matrix,
        "metrics": metrics,
        **metadata,
    }


def _extend_sys_path_with_venv() -> None:
    for site_packages in sorted((PROJECT_ROOT / ".venv" / "lib").glob("python*/site-packages")):
        site_packages_str = str(site_packages)
        if site_packages_str not in sys.path:
            sys.path.append(site_packages_str)
    project_root_str = str(PROJECT_ROOT)
    if project_root_str not in sys.path:
        sys.path.append(project_root_str)


def load_model_module():
    _extend_sys_path_with_venv()
    spec = importlib.util.spec_from_file_location("project_model", PROJECT_ROOT / "model.py")
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load model helpers from {PROJECT_ROOT / 'model.py'}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def row_matches_exclusions(series: pd.Series, exclude_substrings: list[str]) -> pd.Series:
    cleaned = [token.strip().lower() for token in exclude_substrings if token and token.strip()]
    lowered = series.fillna("").astype(str).str.lower()
    mask = pd.Series(False, index=series.index)
    for token in cleaned:
        mask = mask | lowered.str.contains(token, regex=False)
    return mask


def compose_metrics(
    analytical: pd.DataFrame,
    reference_category: str,
    ols_model,
    coefficient_table: pd.DataFrame,
    grouped_importances: pd.DataFrame,
    cv_scores: list[float],
) -> dict[str, float | str]:
    category_importance = grouped_importances.loc[
        grouped_importances["feature_group"] == "category", "importance"
    ]
    top_feature = grouped_importances.iloc[0]["feature_group"]
    second_feature = grouped_importances.iloc[1]["feature_group"] if len(grouped_importances) > 1 else top_feature
    return {
        "n_rows": int(len(analytical)),
        "n_markets": int(analytical["market_id"].nunique()),
        "ols_adj_r2": float(ols_model.rsquared_adj),
        "ols_r2": float(ols_model.rsquared),
        "ols_reference_category": reference_category,
        "rf_mean_r2": float(pd.Series(cv_scores).mean()),
        "rf_std_r2": float(pd.Series(cv_scores).std(ddof=0)),
        "category_importance_share": float(category_importance.iloc[0]) if not category_importance.empty else 0.0,
        "top_feature": str(top_feature),
        "second_feature": str(second_feature),
        "significant_category_terms": int(
            coefficient_table["term"].str.contains("C\\(category").fillna(False)
            .mul(coefficient_table["p_value"] < 0.05)
            .sum()
        ),
    }


def build_filtered_data(exclude_substrings: list[str]) -> dict[str, object]:
    base_analytical = pd.read_csv(DATA_DIR / "analytical_dataset.csv", parse_dates=["date"])
    base_scoring = pd.read_csv(DATA_DIR / "scoring_dataset.csv", parse_dates=["date"])
    analytical_exclusion_mask = row_matches_exclusions(base_analytical["category"], exclude_substrings)
    scoring_exclusion_mask = row_matches_exclusions(base_scoring["category"], exclude_substrings)

    filtered_analytical = base_analytical.loc[~analytical_exclusion_mask].copy()
    filtered_scoring = base_scoring.loc[~scoring_exclusion_mask].copy()
    if filtered_analytical.empty or filtered_scoring.empty:
        raise ValueError("The requested exclusion removed all rows from the report dataset.")

    model_module = load_model_module()
    reference_category = model_module.choose_reference_category(filtered_analytical)
    ols_model, coefficient_table = model_module.fit_ols_model(filtered_analytical, reference_category)
    _, grouped_importances, cv_scores = model_module.fit_random_forest(filtered_analytical)
    latest_scores, category_rankings, volatility_over_time, correlation_matrix = model_module.build_summary_outputs(
        filtered_analytical,
        filtered_scoring,
        ols_model,
        grouped_importances,
    )
    metrics = compose_metrics(
        filtered_analytical,
        reference_category,
        ols_model,
        coefficient_table,
        grouped_importances,
        cv_scores,
    )
    excluded_categories = sorted(base_scoring.loc[scoring_exclusion_mask, "category"].dropna().unique().tolist())
    return {
        "analytical": filtered_analytical,
        "scoring": filtered_scoring,
        "rankings": category_rankings,
        "latest_scores": latest_scores,
        "coefficients": coefficient_table,
        "importances": grouped_importances,
        "volatility_over_time": volatility_over_time,
        "correlation_matrix": correlation_matrix,
        "metrics": metrics,
        "excluded_categories": excluded_categories,
        "excluded_market_count": int(base_scoring.loc[scoring_exclusion_mask, "market_id"].nunique()),
        "excluded_observation_count": int(scoring_exclusion_mask.sum()),
        "filter_label": f"Excluding categories containing: {', '.join(exclude_substrings)}",
        "uses_saved_summary_artifacts": False,
    }


def try_font(name: str, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(name, size)
    except OSError:
        return ImageFont.load_default()


def font_regular(size: int):
    return try_font("DejaVuSans.ttf", size)


def font_bold(size: int):
    return try_font("DejaVuSans-Bold.ttf", size)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def draw_bar_chart(
    labels: list[str],
    values: list[float],
    output_path: Path,
    *,
    title: str,
    value_format: str = "{:.4f}",
) -> None:
    width, height = 1200, 720
    margin_left, margin_right, margin_top, margin_bottom = 220, 60, 110, 70
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font_bold(34)
    axis_font = font_regular(20)
    label_font = font_regular(22)
    value_font = font_regular(18)

    draw.text((margin_left, 35), title, fill="#1F2530", font=title_font)
    chart_top = margin_top
    chart_bottom = height - margin_bottom
    chart_left = margin_left
    chart_right = width - margin_right
    bar_area_height = chart_bottom - chart_top
    bar_height = int(bar_area_height / max(len(labels), 1) * 0.52)
    gap = int(bar_area_height / max(len(labels), 1) * 0.48)
    max_value = max(values) if values else 1

    for idx, (label, value) in enumerate(zip(labels, values)):
        y = chart_top + idx * (bar_height + gap)
        bar_width = 0 if max_value == 0 else int((value / max_value) * (chart_right - chart_left))
        color = CHART_COLORS[idx % len(CHART_COLORS)]
        draw.rounded_rectangle(
            [chart_left, y, chart_left + bar_width, y + bar_height],
            radius=12,
            fill=color,
        )
        draw.text((20, y + 4), label, fill="#1F2530", font=label_font)
        value_text = value_format.format(value)
        draw.text((chart_left + bar_width + 14, y + 6), value_text, fill="#4A5568", font=value_font)

    draw.line([(chart_left, chart_bottom), (chart_right, chart_bottom)], fill="#C9D1DA", width=2)
    draw.text((chart_left, height - 40), "Lower", fill="#7A869A", font=axis_font)
    draw.text((chart_right - 65, height - 40), "Higher", fill="#7A869A", font=axis_font)
    image.save(output_path)


def draw_line_chart(
    time_df: pd.DataFrame,
    output_path: Path,
    *,
    title: str,
) -> None:
    width, height = 1200, 720
    margin_left, margin_right, margin_top, margin_bottom = 100, 60, 110, 90
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font_bold(34)
    axis_font = font_regular(18)
    legend_font = font_regular(18)
    draw.text((margin_left, 35), title, fill="#1F2530", font=title_font)

    plot_left, plot_top = margin_left, margin_top
    plot_right, plot_bottom = width - margin_right, height - margin_bottom
    draw.rectangle([plot_left, plot_top, plot_right, plot_bottom], outline="#D6DEE8", width=2)

    dates = sorted(time_df["date"].unique())
    categories = list(time_df["category"].dropna().unique())
    max_y = max(float(time_df["avg_rolling_volatility_7d"].max()), 0.001)

    for grid_idx in range(5):
        y = plot_bottom - int((plot_bottom - plot_top) * grid_idx / 4)
        draw.line([(plot_left, y), (plot_right, y)], fill="#EEF2F7", width=1)
        value = max_y * grid_idx / 4
        draw.text((18, y - 10), f"{value:.3f}", fill="#7A869A", font=axis_font)

    if len(dates) > 1:
        total_days = max((dates[-1] - dates[0]).days, 1)
    else:
        total_days = 1

    for idx, category in enumerate(categories):
        subset = time_df[time_df["category"] == category].sort_values("date")
        color = CHART_COLORS[idx % len(CHART_COLORS)]
        points = []
        for row in subset.itertuples():
            day_offset = (row.date - dates[0]).days if len(dates) > 1 else 0
            x = plot_left + int((plot_right - plot_left) * day_offset / total_days)
            y = plot_bottom - int((plot_bottom - plot_top) * (row.avg_rolling_volatility_7d / max_y))
            points.append((x, y))
        if len(points) >= 2:
            draw.line(points, fill=color, width=4)
        elif points:
            x, y = points[0]
            draw.ellipse([x - 3, y - 3, x + 3, y + 3], fill=color)

    tick_positions = [0.0, 0.33, 0.66, 1.0]
    for position in tick_positions:
        x = plot_left + int((plot_right - plot_left) * position)
        draw.line([(x, plot_bottom), (x, plot_bottom + 8)], fill="#C9D1DA", width=2)
        tick_date = dates[0] + (dates[-1] - dates[0]) * position if len(dates) > 1 else dates[0]
        draw.text((x - 35, plot_bottom + 18), tick_date.strftime("%b %Y"), fill="#7A869A", font=axis_font)

    legend_y = height - 46
    legend_x = plot_left
    for idx, category in enumerate(categories):
        color = CHART_COLORS[idx % len(CHART_COLORS)]
        draw.rounded_rectangle([legend_x, legend_y, legend_x + 22, legend_y + 10], radius=4, fill=color)
        draw.text((legend_x + 30, legend_y - 7), category, fill="#1F2530", font=legend_font)
        legend_x += 165
        if legend_x > plot_right - 150:
            legend_x = plot_left
            legend_y -= 28

    image.save(output_path)


def draw_heatmap(corr_df: pd.DataFrame, output_path: Path, *, title: str) -> None:
    width, height = 1200, 760
    margin_left, margin_right, margin_top, margin_bottom = 240, 70, 120, 120
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font_bold(34)
    label_font = font_regular(18)
    cell_font = font_regular(17)
    draw.text((margin_left, 35), title, fill="#1F2530", font=title_font)

    rows = list(corr_df.index)
    cols = list(corr_df.columns)
    grid_w = width - margin_left - margin_right
    grid_h = height - margin_top - margin_bottom
    cell_w = grid_w / max(len(cols), 1)
    cell_h = grid_h / max(len(rows), 1)

    def cell_color(value: float) -> tuple[int, int, int]:
        value = max(-1.0, min(1.0, float(value)))
        if value >= 0:
            base = int(255 - value * 120)
            return (base, base, 255)
        base = int(255 - abs(value) * 120)
        return (255, base, base)

    for i, row_name in enumerate(rows):
        y0 = margin_top + i * cell_h
        draw.text((20, y0 + cell_h / 2 - 8), row_name.replace("_", " "), fill="#1F2530", font=label_font)
        for j, col_name in enumerate(cols):
            x0 = margin_left + j * cell_w
            if i == 0:
                text_box = Image.new("RGBA", (180, 38), (255, 255, 255, 0))
                text_draw = ImageDraw.Draw(text_box)
                text_draw.text((0, 8), col_name.replace("_", " "), fill="#1F2530", font=label_font)
                rotated = text_box.rotate(35, expand=1, resample=Image.Resampling.BICUBIC)
                image.paste(rotated, (int(x0 + 8), 70), rotated)
            value = float(corr_df.iloc[i, j])
            draw.rectangle(
                [x0, y0, x0 + cell_w, y0 + cell_h],
                fill=cell_color(value),
                outline="#D8DEE9",
            )
            text = f"{value:.2f}"
            bbox = draw.textbbox((0, 0), text, font=cell_font)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            draw.text(
                (x0 + cell_w / 2 - tw / 2, y0 + cell_h / 2 - th / 2),
                text,
                fill="#1F2530",
                font=cell_font,
            )

    image.save(output_path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = GRID_GRAY) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, *, name: str = "Calibri", size: int | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_body_paragraph(doc: Document, text: str, *, after: int = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=BODY_COLOR)


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, color=ACCENT_BLUE if level == 1 else ACCENT_DARK, bold=True)


def add_title_page(
    doc: Document,
    metrics: dict,
    scoring: pd.DataFrame,
    *,
    report_title: str,
    subtitle: str,
    scope_note: str,
    extra_meta_rows: list[tuple[str, str]] | None = None,
) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Prediction Market Volatility Modeling")
    set_run_font(header_run, size=9, color=MUTED_COLOR)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Outputs report generated from committed MVP artifacts")
    set_run_font(footer_run, size=9, color=MUTED_COLOR)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(report_title)
    set_run_font(title_run, size=24, color=BODY_COLOR, bold=True)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    set_run_font(subtitle_run, size=14, color=MUTED_COLOR)

    meta_rows = [
        ("Research question", "Can market category predict future volatility after controls?"),
        ("Generated", date.today().isoformat()),
        ("Markets", f"{metrics['n_markets']:,}"),
        ("Training rows", f"{metrics['n_rows']:,}"),
        ("Date span", f"{scoring['date'].min().date()} to {scoring['date'].max().date()}"),
        ("Reference category", str(metrics["ols_reference_category"])),
    ]
    if extra_meta_rows:
        meta_rows.extend(extra_meta_rows)
    add_key_value_table(doc, meta_rows, widths_dxa=[1800, 7560])

    add_body_paragraph(
        doc,
        scope_note,
        after=10,
    )


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        row = table.add_row()
        label_cell, value_cell = row.cells
        label_cell.text = ""
        value_cell.text = ""
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(label)
        set_run_font(lr, size=10, color=MUTED_COLOR, bold=True)
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(value)
        set_run_font(vr, size=10, color=BODY_COLOR)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(label_cell)
        set_cell_border(value_cell)
        set_cell_shading(label_cell, HEADER_FILL)
    TABLE_GEOMETRY.apply_table_geometry(
        table,
        TABLE_GEOMETRY.exact_column_widths(widths_dxa, CONTENT_WIDTH_DXA),
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )


def add_dataframe_table(
    doc: Document,
    df: pd.DataFrame,
    widths_dxa: list[int],
    *,
    column_labels: list[str] | None = None,
    numeric_cols: set[str] | None = None,
) -> None:
    numeric_cols = numeric_cols or set()
    column_labels = column_labels or list(df.columns)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_cells = table.rows[0].cells
    for idx, label in enumerate(column_labels):
        cell = header_cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(label))
        set_run_font(r, size=9, color=BODY_COLOR, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)

    for row_data in df.itertuples(index=False):
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if df.columns[idx] in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            text = str(value)
            r = p.add_run(text)
            set_run_font(r, size=9, color=BODY_COLOR)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)

    TABLE_GEOMETRY.apply_table_geometry(
        table,
        TABLE_GEOMETRY.exact_column_widths(widths_dxa, CONTENT_WIDTH_DXA),
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )


def add_chart(doc: Document, image_path: Path, caption: str) -> None:
    doc.add_picture(str(image_path), width=Inches(6.05))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(10)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=MUTED_COLOR)


def fit_term(term: str) -> str:
    replacements = {
        "C(category, Treatment(reference='US Elections'))[T.": "",
        "]": "",
        "recent_volatility_7d": "Recent volatility (7d)",
        "bid_ask_spread": "Bid-ask spread",
        "log_volume": "Log volume",
        "prob_level": "Probability level",
        "time_to_expiry_days": "Time to expiry (days)",
        "Intercept": "Intercept",
    }
    result = term
    for source, target in replacements.items():
        result = result.replace(source, target)
    return result


def summarize_category_effects(coefficients: pd.DataFrame, reference_category: str) -> str:
    category_coefficients = coefficients[
        coefficients["term"].str.contains("C(category", regex=False, na=False)
    ].copy()
    if category_coefficients.empty:
        return f"No category-dummy terms remain after filtering; {reference_category} is the only category in scope."

    highest = category_coefficients.sort_values("coef", ascending=False).iloc[0]
    lowest = category_coefficients.sort_values("coef", ascending=True).iloc[0]
    highest_label = fit_term(highest["term"])
    lowest_label = fit_term(lowest["term"])
    return (
        f"Relative to {reference_category}, the strongest positive category effect is {highest_label} "
        f"({highest['coef']:.4f}, p={highest['p_value']:.3g}), while the lowest category effect is "
        f"{lowest_label} ({lowest['coef']:.4f}, p={lowest['p_value']:.3g})."
    )


def format_volume(value: float) -> str:
    value = float(value)
    if abs(value) >= 1_000_000_000:
        return f"{value / 1_000_000_000:.1f}B"
    if abs(value) >= 1_000_000:
        return f"{value / 1_000_000:.1f}M"
    if abs(value) >= 1_000:
        return f"{value / 1_000:.1f}K"
    return f"{value:,.0f}"


def build_document(
    *,
    data_dir: Path = DATA_DIR,
    exclude_substrings: list[str] | None = None,
    output_docx: Path | None = None,
) -> Path:
    ensure_dirs()
    exclude_substrings = [token for token in (exclude_substrings or []) if token.strip()]
    data = build_filtered_data(exclude_substrings) if exclude_substrings else load_saved_data(data_dir)
    analytical = data["analytical"]
    scoring = data["scoring"]
    rankings = data["rankings"]
    latest_scores = data["latest_scores"]
    coefficients = data["coefficients"]
    importances = data["importances"]
    volatility_over_time = data["volatility_over_time"]
    correlation_matrix = data["correlation_matrix"]
    metrics = data["metrics"]
    excluded_categories = data["excluded_categories"]
    excluded_market_count = data["excluded_market_count"]
    excluded_observation_count = data["excluded_observation_count"]
    uses_saved_summary_artifacts = bool(data["uses_saved_summary_artifacts"])

    if exclude_substrings:
        asset_suffix = "excluding_" + "_".join(token.lower().replace(" ", "_") for token in exclude_substrings)
    elif data_dir != DATA_DIR:
        asset_suffix = data_dir.name.lower().replace(" ", "_")
    else:
        asset_suffix = "full_sample"
    asset_dir = DEFAULT_ASSET_DIR if data_dir == DATA_DIR and not exclude_substrings else DELIVERABLE_DIR / f"outputs_report_assets_{asset_suffix}"
    asset_dir.mkdir(parents=True, exist_ok=True)
    if output_docx is None:
        output_docx = (
            DEFAULT_OUTPUT_DOCX
            if data_dir == DATA_DIR and not exclude_substrings
            else DELIVERABLE_DIR / f"prediction_market_volatility_outputs_report_{asset_suffix}.docx"
        )

    category_chart = asset_dir / "category_rankings.png"
    importance_chart = asset_dir / "feature_importance.png"
    line_chart = asset_dir / "volatility_over_time.png"
    heatmap_chart = asset_dir / "correlation_heatmap.png"
    report_title = "Prediction Market Volatility Outputs Report"
    subtitle = "Polymarket category-volatility MVP"
    scope_note = (
        "This document consolidates the current committed project outputs into one report: data coverage, "
        "category rankings, model metrics, coefficient tables, feature importances, high-risk markets, "
        "and the saved artifacts consumed by the notebook and Streamlit dashboard."
    )
    extra_meta_rows: list[tuple[str, str]] = []
    if exclude_substrings or excluded_categories:
        report_title = "Prediction Market Volatility Outputs Report (Crypto Excluded)"
        subtitle = "Polymarket category-volatility MVP | non-crypto slice"
        scope_note = (
            "This document recomputes the analysis after excluding categories whose labels contain the requested "
            f"token(s): {', '.join(exclude_substrings or excluded_categories)}. All tables, charts, OLS coefficients, random-forest "
            "importances, and market-risk rankings in this report are derived from the filtered non-crypto subset."
        )
        extra_meta_rows = [
            ("Excluded categories", ", ".join(excluded_categories) if excluded_categories else ", ".join(exclude_substrings)),
            ("Excluded markets", f"{excluded_market_count:,}"),
            ("Excluded observations", f"{excluded_observation_count:,}"),
        ]

    draw_bar_chart(
        rankings["category"].tolist(),
        rankings["median_volatility"].tolist(),
        category_chart,
        title="Median 7-day volatility by category",
        value_format="{:.4f}",
    )
    draw_bar_chart(
        importances["feature_group"].str.replace("_", " ").str.title().tolist(),
        importances["importance"].tolist(),
        importance_chart,
        title="Random forest feature importance share",
        value_format="{:.3f}",
    )
    draw_line_chart(
        volatility_over_time.sort_values("date"),
        line_chart,
        title="Average category-level volatility over time",
    )
    draw_heatmap(
        correlation_matrix,
        heatmap_chart,
        title="Analytical feature correlation matrix",
    )

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_COLOR
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.1

    for name, size, color in [("Heading 1", 16, ACCENT_BLUE), ("Heading 2", 13, ACCENT_BLUE), ("Heading 3", 12, ACCENT_DARK)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    add_title_page(
        doc,
        metrics,
        scoring,
        report_title=report_title,
        subtitle=subtitle,
        scope_note=scope_note,
        extra_meta_rows=extra_meta_rows,
    )

    add_section_heading(doc, "Executive Summary", level=1)
    key_findings = [
        (
            f"The current sample contains {metrics['n_markets']:,} markets and {metrics['n_rows']:,} "
            f"training rows spanning {scoring['date'].min().date()} to {scoring['date'].max().date()}."
        ),
        (
            f"The OLS model reaches an adjusted R² of {metrics['ols_adj_r2']:.3f}. "
            f"The random forest reaches a mean time-series cross-validation R² of {metrics['rf_mean_r2']:.3f}."
        ),
        (summarize_category_effects(coefficients, str(metrics["ols_reference_category"]))),
        (
            f"The random forest ranks {str(metrics['top_feature']).replace('_', ' ')} first and "
            f"{str(metrics['second_feature']).replace('_', ' ')} second. Category contributes "
            f"{metrics['category_importance_share']:.1%} of total feature importance."
        ),
    ]
    if exclude_substrings or excluded_categories:
        key_findings.insert(
            1,
            (
                f"This slice excludes {excluded_market_count:,} markets across the category bucket(s) "
                f"{', '.join(excluded_categories) if excluded_categories else ', '.join(exclude_substrings)}, "
                f"removing {excluded_observation_count:,} daily scoring observations from the report scope."
            ),
        )
    for paragraph in key_findings:
        add_body_paragraph(doc, paragraph)

    add_section_heading(doc, "Coverage Snapshot", level=1)
    snapshot_rows = [
        ("Markets", f"{metrics['n_markets']:,}"),
        ("Analytical rows", f"{metrics['n_rows']:,}"),
        ("Date span", f"{scoring['date'].min().date()} to {scoring['date'].max().date()}"),
        ("Reference category", str(metrics["ols_reference_category"])),
        ("Top feature", str(metrics["top_feature"]).replace("_", " ").title()),
        ("Second feature", str(metrics["second_feature"]).replace("_", " ").title()),
        ("Significant category terms", f"{metrics['significant_category_terms']}"),
        ("Category importance share", f"{metrics['category_importance_share']:.1%}"),
    ]
    if exclude_substrings or excluded_categories:
        snapshot_rows.extend(
            [
                ("Excluded categories", ", ".join(excluded_categories) if excluded_categories else ", ".join(exclude_substrings)),
                ("Excluded markets", f"{excluded_market_count:,}"),
            ]
        )
    add_key_value_table(doc, snapshot_rows, widths_dxa=[2400, 6960])

    add_section_heading(doc, "Category Rankings Output", level=1)
    add_body_paragraph(
        doc,
        "The ranking chart and table below summarize realized 7-day volatility by expanded market category.",
    )
    add_chart(doc, category_chart, "Figure 1. Median realized 7-day volatility by expanded category.")
    ranking_table = rankings.copy()
    ranking_table["median_volatility"] = ranking_table["median_volatility"].map(lambda x: f"{x:.4f}")
    ranking_table["mean_volatility"] = ranking_table["mean_volatility"].map(lambda x: f"{x:.4f}")
    ranking_table["avg_volume"] = ranking_table["avg_volume"].map(format_volume)
    add_dataframe_table(
        doc,
        ranking_table[["category", "market_count", "observation_count", "median_volatility", "mean_volatility", "avg_volume"]],
        widths_dxa=[2100, 1080, 1380, 1260, 1260, 2280],
        column_labels=["Category", "Markets", "Observations", "Median 7d vol", "Mean 7d vol", "Avg volume"],
        numeric_cols={"market_count", "observation_count", "median_volatility", "mean_volatility"},
    )

    add_section_heading(doc, "Time-Series and Correlation Outputs", level=1)
    add_body_paragraph(
        doc,
        "These charts show how average category volatility evolves over time and how the engineered predictors move together in the analytical panel.",
    )
    add_chart(doc, line_chart, "Figure 2. Average category-level volatility over calendar time.")
    add_chart(doc, heatmap_chart, "Figure 3. Correlation matrix for core analytical variables.")

    add_section_heading(doc, "OLS Regression Output", level=1)
    add_body_paragraph(
        doc,
        "The coefficient tables below separate category effects from control-variable effects so the category signal is easier to read.",
    )
    category_coefficients = coefficients[coefficients["term"].str.contains("C(category", regex=False, na=False)].copy()
    category_coefficients["term"] = category_coefficients["term"].map(fit_term)
    category_coefficients["coef"] = category_coefficients["coef"].map(lambda x: f"{x:.4f}")
    category_coefficients["p_value"] = category_coefficients["p_value"].map(lambda x: f"{x:.3g}")
    add_dataframe_table(
        doc,
        category_coefficients[["term", "coef", "p_value"]],
        widths_dxa=[5760, 1800, 1800],
        column_labels=["Category term vs reference", "Coefficient", "P-value"],
        numeric_cols={"coef", "p_value"},
    )

    control_coefficients = coefficients[~coefficients["term"].str.contains("C(category", regex=False, na=False)].copy()
    control_coefficients["term"] = control_coefficients["term"].map(fit_term)
    control_coefficients["coef"] = control_coefficients["coef"].map(lambda x: f"{x:.4f}")
    control_coefficients["p_value"] = control_coefficients["p_value"].map(lambda x: f"{x:.3g}")
    add_body_paragraph(doc, "")
    add_dataframe_table(
        doc,
        control_coefficients[["term", "coef", "p_value"]],
        widths_dxa=[5760, 1800, 1800],
        column_labels=["Control term", "Coefficient", "P-value"],
        numeric_cols={"coef", "p_value"},
    )

    add_section_heading(doc, "Random Forest Output", level=1)
    add_body_paragraph(
        doc,
        "The random forest captures non-linear predictive structure and provides a separate ranking of feature importance.",
    )
    add_chart(doc, importance_chart, "Figure 4. Aggregated random-forest feature importance share.")
    importance_table = importances.copy()
    importance_table["feature_group"] = importance_table["feature_group"].str.replace("_", " ").str.title()
    importance_table["importance"] = importance_table["importance"].map(lambda x: f"{x:.3f}")
    add_dataframe_table(
        doc,
        importance_table,
        widths_dxa=[5760, 3600],
        column_labels=["Feature group", "Importance share"],
        numeric_cols={"importance"},
    )

    add_section_heading(doc, "High-Risk Market Output", level=1)
    add_body_paragraph(
        doc,
        "The table below shows the current top-risk markets based on the normalized OLS-predicted forward-volatility score.",
    )
    top_risk = latest_scores.sort_values("risk_score", ascending=False).head(12).copy()
    top_risk["implied_probability"] = top_risk["implied_probability"].map(lambda x: f"{x:.3f}")
    top_risk["risk_score"] = top_risk["risk_score"].map(lambda x: f"{x:.2f}")
    top_risk["time_to_expiry_days"] = top_risk["time_to_expiry_days"].map(lambda x: f"{x:.0f}")
    top_risk["recent_volatility_7d"] = top_risk["recent_volatility_7d"].map(lambda x: f"{x:.4f}")
    top_risk["title"] = top_risk["title"].map(lambda x: "\n".join(textwrap.wrap(x, width=42)))
    add_dataframe_table(
        doc,
        top_risk[["title", "category", "risk_score", "implied_probability", "time_to_expiry_days", "recent_volatility_7d"]],
        widths_dxa=[3300, 1440, 900, 1080, 1200, 1440],
        column_labels=["Market title", "Category", "Risk score", "Probability", "Days to expiry", "Current 7d vol"],
        numeric_cols={"risk_score", "implied_probability", "time_to_expiry_days", "recent_volatility_7d"},
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_section_heading(doc, "Saved Artifact Inventory", level=1)
    add_body_paragraph(
        doc,
        (
            "This appendix lists the saved outputs backing the notebook and Streamlit application."
            if uses_saved_summary_artifacts
            else "This appendix lists the source artifacts used to recompute the filtered non-crypto report."
        ),
    )
    if uses_saved_summary_artifacts:
        artifact_rows = pd.DataFrame(
            [
                ("raw_markets.csv", "Market-level metadata pulled from Polymarket"),
                ("raw_prices.csv", "Raw price-history observations"),
                ("analytical_dataset.csv", "Clean modeling dataset with forward-volatility target"),
                ("scoring_dataset.csv", "Scoring dataset used for dashboard market snapshots"),
                ("latest_market_scores.csv", "Latest market-level risk scores"),
                ("category_rankings.csv", "Category-level volatility summary"),
                ("volatility_over_time.csv", "Date x category average volatility series"),
                ("correlation_matrix.csv", "Analytical correlation matrix"),
                ("ols_coefficients.csv", "OLS regression coefficient table"),
                ("rf_feature_importances.csv", "Aggregated random-forest importances"),
                ("model_metrics.json", "Summary metrics consumed by the dashboard"),
                ("notebooks/analysis.ipynb", "Notebook version of the analysis"),
            ],
            columns=["Artifact", "Purpose"],
        )
    else:
        artifact_rows = pd.DataFrame(
            [
                ("raw_markets.csv", "Underlying market metadata source"),
                ("raw_prices.csv", "Underlying raw price-history source"),
                ("analytical_dataset.csv", "Source analytical dataset filtered in-memory to exclude crypto categories"),
                ("scoring_dataset.csv", "Source scoring dataset filtered in-memory to exclude crypto categories"),
                ("build_outputs_document.py", "Report builder used to recompute outputs on the filtered slice"),
                (output_docx.name, "Final non-crypto document deliverable"),
            ],
            columns=["Artifact", "Purpose"],
        )
    add_dataframe_table(
        doc,
        artifact_rows,
        widths_dxa=[2640, 6720],
        column_labels=["Artifact", "Purpose"],
    )

    doc.save(output_docx)
    return output_docx


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build a DOCX report for prediction-market volatility outputs.")
    parser.add_argument("--data-dir", type=Path, default=DATA_DIR)
    parser.add_argument("--exclude-substring", action="append", default=[], help="Exclude categories containing this substring.")
    parser.add_argument("--output-docx", type=Path, default=None)
    return parser


if __name__ == "__main__":
    args = build_arg_parser().parse_args()
    path = build_document(data_dir=args.data_dir, exclude_substrings=args.exclude_substring, output_docx=args.output_docx)
    print(path)
